import asyncio
import io
import json
import os
import re
import time
from typing import Any, Dict, List, Optional

import httpx
import structlog
from nicegui import events, ui
from utils.api_client import RAESuiteClient

logger = structlog.get_logger(__name__)

# Endpoints configuration (Kubernetes In-Cluster Services with fallback)
NODE1_OLLAMA_URL = os.getenv("NODE1_OLLAMA_URL", "http://ollama-node1:11434")
LAPTOP_OLLAMA_URL = os.getenv("LAPTOP_OLLAMA_URL", "http://ollama-laptop:11434")
CLOUD_OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")

AVAILABLE_MODELS = {
    "node1_bielik": {
        "name": "Bielik 11B v3 (Node 1 - RTX 4080)",
        "short_name": "Bielik 11B v3",
        "device": "Node 1 (Lumina RTX 4080)",
        "badge_color": "indigo",
        "icon": "rocket_launch",
        "type": "ollama",
        "url": NODE1_OLLAMA_URL,
        "model": "SpeakLeash/bielik-11b-v3.0-instruct:Q5_K_M",
        "desc": "Polski model językowy SpeakLeash na karcie NVIDIA RTX 4080 SUPER (16 GB VRAM)",
    },
    "node1_qwen35": {
        "name": "Qwen 3.5 9B (Node 1 - RTX 4080)",
        "short_name": "Qwen 3.5 9B",
        "device": "Node 1 (Lumina RTX 4080)",
        "badge_color": "indigo",
        "icon": "psychology",
        "type": "ollama",
        "url": NODE1_OLLAMA_URL,
        "model": "qwen3.5:9b",
        "desc": "Szybki, precyzyjny model ogólny z wbudowanym wnioskowaniem (Node 1)",
    },
    "node1_deepseek_r1": {
        "name": "DeepSeek-R1 8B (Node 1 - RTX 4080)",
        "short_name": "DeepSeek-R1 8B",
        "device": "Node 1 (Lumina RTX 4080)",
        "badge_color": "purple",
        "icon": "insights",
        "type": "ollama",
        "url": NODE1_OLLAMA_URL,
        "model": "deepseek-r1:8b",
        "desc": "Zaawansowany model rozumowania logicznego i audytu umów (Node 1)",
    },
    "node1_gemma4": {
        "name": "Gemma 4 12B (Node 1 - RTX 4080)",
        "short_name": "Gemma 4 12B",
        "device": "Node 1 (Lumina RTX 4080)",
        "badge_color": "blue",
        "icon": "hub",
        "type": "ollama",
        "url": NODE1_OLLAMA_URL,
        "model": "gemma4:12b",
        "desc": "Model Google Gemma 4 na stacji Node 1",
    },
    "laptop_bielik": {
        "name": "Bielik 11B v3 (Laptop GPU)",
        "short_name": "Bielik 11B (Laptop)",
        "device": "Laptop (Lokalne GPU)",
        "badge_color": "teal",
        "icon": "laptop_chromebook",
        "type": "ollama",
        "url": LAPTOP_OLLAMA_URL,
        "model": "SpeakLeash/bielik-11b-v3.0-instruct:Q5_K_M",
        "desc": "Polski model Bielik uruchomiony lokalnie na laptopie",
    },
    "laptop_qwen35": {
        "name": "Qwen 3.5 9B (Laptop GPU)",
        "short_name": "Qwen 3.5 (Laptop)",
        "device": "Laptop (Lokalne GPU)",
        "badge_color": "teal",
        "icon": "laptop_chromebook",
        "type": "ollama",
        "url": LAPTOP_OLLAMA_URL,
        "model": "qwen3.5:9b",
        "desc": "Model Qwen 3.5 uruchomiony lokalnie na laptopie",
    },
    "laptop_deepseek_r1": {
        "name": "DeepSeek-R1 8B (Laptop GPU)",
        "short_name": "DeepSeek-R1 (Laptop)",
        "device": "Laptop (Lokalne GPU)",
        "badge_color": "teal",
        "icon": "laptop_chromebook",
        "type": "ollama",
        "url": LAPTOP_OLLAMA_URL,
        "model": "deepseek-r1:8b",
        "desc": "Model DeepSeek-R1 uruchomiony lokalnie na laptopie",
    },
    "laptop_gemma4": {
        "name": "Gemma 4 12B (Laptop GPU)",
        "short_name": "Gemma 4 (Laptop)",
        "device": "Laptop (Lokalne GPU)",
        "badge_color": "teal",
        "icon": "laptop_chromebook",
        "type": "ollama",
        "url": LAPTOP_OLLAMA_URL,
        "model": "gemma4:12b",
        "desc": "Model Gemma 4 uruchomiony lokalnie na laptopie",
    },
    "cloud_openrouter": {
        "name": "Chmura OpenRouter (Gemini 3.7 Flash)",
        "short_name": "Gemini 3.7 Flash (Cloud)",
        "device": "Chmura Zewnętrzna",
        "badge_color": "amber",
        "icon": "cloud",
        "type": "openrouter",
        "url": CLOUD_OPENROUTER_URL,
        "model": "google/gemini-3.7-flash",
        "desc": "Szybki model chmurowy Google Gemini 3.7 Flash przez OpenRouter",
    },
}


class PDKnowledgeHubApp:
    """
    Print & Display Corporate Document Analyzer & AI Knowledge Hub.
    Scoped to RAE Project: 'pd-test-llm'.
    """

    def __init__(self, client: RAESuiteClient, user_email: str = "user@printdisplay.pl"):
        self.client = client
        self.project = "pd-test-llm"
        self.user_email = user_email
        self.messages: List[Dict[str, Any]] = []
        self.chat_container = None
        self.docs_badge_label = None
        self.docs_sidebar = None
        self.is_generating = False
        self.selected_model_key = "node1_bielik"
        self.use_rag = True
        self.use_legal_audit = True
        self.use_web_search = True
        self.is_dark_mode = False
        self.dark_mode_controller = None
        self.model_select_widget = None
        self.active_model_badge = None
        self.prompt_input = None

    # --- Live Web & Legal Search (ISAP Sejm + Wikipedia API) ---

    @staticmethod
    async def search_isap_and_web(query: str) -> List[Dict[str, str]]:
        """
        Searches the official Polish Sejm ISAP API and Legal Knowledge bases.
        """
        results = []
        headers = {
            "Accept": "application/json",
            "User-Agent": "PDKnowledgeHub/1.0 (legal@printdisplay.pl)",
        }

        # 1. Search ISAP Sejm API
        clean_terms = re.sub(r"[^\w\s]", "", query).split()
        search_term = clean_terms[0] if clean_terms else "kodeks"
        for t in ["cywilny", "autorskim", "umowa", "zapłaty", "rodo", "pracy"]:
            if t in query.lower():
                search_term = t
                break

        try:
            async with httpx.AsyncClient(timeout=8.0) as client:
                isap_url = f"https://api.sejm.gov.pl/eli/acts/search?title={search_term}&limit=4"
                r = await client.get(isap_url, headers=headers)
                if r.status_code == 200:
                    items = r.json().get("items", [])
                    for item in items[:3]:
                        results.append({
                            "source": f"ISAP Sejm RP ({item.get('ELI', 'D.U.')})",
                            "title": item.get("title", ""),
                            "snippet": f"Status aktu: {item.get('status', 'obowiązujący')}. Data ogłoszenia: {item.get('promulgation', '-')}",
                            "url": f"https://isap.sejm.gov.pl/isap.nsf/DocDetails.xsp?id={item.get('ELI', '').replace('/', '')}",
                        })
        except Exception as e:
            logger.warning("isap_search_failed", error=str(e))

        # 2. Search Legal Wiki/Doctrine
        try:
            async with httpx.AsyncClient(timeout=8.0) as client:
                wiki_url = f"https://pl.wikipedia.org/w/api.php?action=query&list=search&srsearch={query}&format=json"
                r = await client.get(wiki_url, headers=headers)
                if r.status_code == 200:
                    data = r.json()
                    wiki_items = data.get("query", {}).get("search", [])
                    for item in wiki_items[:2]:
                        snippet = re.sub(r"<[^<]+?>", "", item.get("snippet", ""))
                        results.append({
                            "source": f"Encyklopedia Prawna / {item.get('title')}",
                            "title": item.get("title", ""),
                            "snippet": snippet,
                            "url": f"https://pl.wikipedia.org/wiki/{item.get('title', '').replace(' ', '_')}",
                        })
        except Exception as e:
            logger.warning("wiki_legal_search_failed", error=str(e))

        return results

    # --- Document Parsing Engines ---

    @staticmethod
    def _extract_text_from_pdf(content_bytes: bytes) -> str:
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            text_pages = []
            for i, page in enumerate(reader.pages):
                txt = page.extract_text()
                if txt:
                    text_pages.append(f"--- [Strona {i+1}] ---\n{txt.strip()}")
            return "\n\n".join(text_pages)
        except Exception as e:
            logger.error("pdf_extract_failed", error=str(e))
            return content_bytes.decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_text_from_docx(content_bytes: bytes) -> str:
        try:
            import docx

            doc = docx.Document(io.BytesIO(content_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_txt:
                        paragraphs.append(row_txt)
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error("docx_extract_failed", error=str(e))
            return content_bytes.decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_text_from_xlsx(content_bytes: bytes) -> str:
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
            lines = []
            for sheetname in wb.sheetnames:
                sheet = wb[sheetname]
                lines.append(f"=== Arkusz: {sheetname} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [
                        str(v).strip()
                        for v in row
                        if v is not None and str(v).strip()
                    ]
                    if row_vals:
                        lines.append(" | ".join(row_vals))
            return "\n".join(lines)
        except Exception as e:
            logger.error("xlsx_extract_failed", error=str(e))
            return content_bytes.decode("utf-8", errors="ignore")

    @classmethod
    def extract_document_text(cls, filename: str, content_bytes: bytes) -> str:
        ext = filename.lower().split(".")[-1]
        if ext == "pdf":
            return cls._extract_text_from_pdf(content_bytes)
        elif ext in ["docx", "doc"]:
            return cls._extract_text_from_docx(content_bytes)
        elif ext in ["xlsx", "xls"]:
            return cls._extract_text_from_xlsx(content_bytes)
        else:
            return content_bytes.decode("utf-8", errors="ignore")

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 80) -> List[str]:
        words = re.split(r"\s+", text)
        chunks = []
        i = 0
        while i < len(words):
            chunk = " ".join(words[i : i + chunk_size])
            if chunk.strip():
                chunks.append(chunk.strip())
            i += chunk_size - overlap
        return chunks if chunks else [text]

    # --- Document Ingestion & Storage ---

    async def handle_file_upload(self, e: events.UploadEventArguments):
        filename = e.name
        content_bytes = e.content.read()
        size_kb = round(len(content_bytes) / 1024, 1)

        ui.notify(
            f"📄 Rozpoczynam przetwarzanie: {filename} ({size_kb} KB)...",
            type="info",
            position="top-right",
        )

        try:
            extracted_text = await asyncio.to_thread(
                self.extract_document_text, filename, content_bytes
            )

            if not extracted_text or len(extracted_text.strip()) < 10:
                ui.notify(
                    f"⚠️ Nie udało się wyodrębnić tekstu z pliku: {filename}",
                    type="warning",
                    position="top-right",
                )
                return

            chunks = self.chunk_text(extracted_text, chunk_size=500, overlap=80)
            total_chunks = len(chunks)

            ui.notify(
                f"🧠 Indeksowanie {total_chunks} fragmentów w RAE Memory...",
                type="info",
                position="top-right",
            )

            stored_count = 0
            for idx, chunk in enumerate(chunks):
                metadata = {
                    "filename": filename,
                    "chunk_index": idx + 1,
                    "total_chunks": total_chunks,
                    "uploaded_by": self.user_email,
                    "doc_title": filename,
                    "file_size_kb": size_kb,
                    "char_count": len(chunk),
                }
                success = await self.client.store_memory(
                    content=chunk,
                    project=self.project,
                    source=f"upload:{filename}",
                    tags=["pd_knowledge", "document", filename.split(".")[-1]],
                    metadata=metadata,
                    layer="semantic",
                    importance=0.85,
                )
                if success:
                    stored_count += 1

            ui.notify(
                f"✅ Zindeksowano {stored_count}/{total_chunks} fragmentów z '{filename}' w bazie wiedzy Print & Display!",
                type="positive",
                position="top-right",
                duration=5.0,
            )

            # Insert a notice into chat
            if self.chat_container:
                with self.chat_container:
                    with ui.row().classes(
                        "w-full justify-center my-2"
                    ):
                        with ui.card().classes(
                            "bg-blue-50 dark:bg-slate-800 border border-blue-200 dark:border-blue-700/50 px-4 py-2 rounded-full"
                        ):
                            ui.label(
                                f"📥 Dodano dokument: {filename} ({stored_count} fragmentów)"
                            ).classes(
                                "text-xs font-semibold text-blue-700 dark:text-blue-300"
                            )

            await self.refresh_documents_list()

        except Exception as err:
            logger.error("upload_failed", filename=filename, error=str(err))
            ui.notify(
                f"❌ Błąd wgrywania: {str(err)}",
                type="negative",
                position="top-right",
            )

    async def delete_document_chunks(self, filename: str):
        ui.notify(
            f"Usuwanie dokumentu '{filename}'...",
            type="info",
            position="top-right",
        )
        try:
            memories_resp = await self.client.list_memories(
                project=self.project, limit=500
            )
            items = memories_resp.get("results", [])
            deleted = 0
            for item in items:
                source = item.get("source") or ""
                meta_fn = item.get("metadata", {}).get("filename", "")
                if source == f"upload:{filename}" or meta_fn == filename:
                    await self.client.delete_memory(item["id"])
                    deleted += 1

            ui.notify(
                f"Usunięto {deleted} fragmentów pliku '{filename}'.",
                type="positive",
                position="top-right",
            )
            await self.refresh_documents_list()
        except Exception as e:
            ui.notify(
                f"Błąd podczas usuwania: {str(e)}",
                type="negative",
                position="top-right",
            )

    async def refresh_documents_list(self):
        if not self.docs_sidebar:
            return
        self.docs_sidebar.clear()

        try:
            memories_resp = await self.client.list_memories(
                project=self.project, limit=500
            )
            items = memories_resp.get("results", [])

            docs_map: Dict[str, Dict[str, Any]] = {}
            for item in items:
                meta = item.get("metadata", {})
                fn = meta.get("filename") or item.get("source", "Inne").replace(
                    "upload:", ""
                )
                if fn not in docs_map:
                    is_legal = meta.get("is_legal_reference", False) or "Kodeks" in fn or "Prawo" in fn or "Ustawa" in fn or "RODO" in fn or "Standardy" in fn
                    docs_map[fn] = {
                        "filename": fn,
                        "chunks": 0,
                        "uploaded_by": meta.get("uploaded_by", "system"),
                        "size_kb": meta.get("file_size_kb", "-"),
                        "sample_text": item.get("content", "")[:150],
                        "is_legal": is_legal,
                    }
                docs_map[fn]["chunks"] += 1

            if self.docs_badge_label:
                self.docs_badge_label.text = f"{len(docs_map)} dok."

            with self.docs_sidebar:
                if not docs_map:
                    with ui.column().classes("w-full items-center p-6 text-center"):
                        ui.icon("folder_open", size="3rem", color="slate-400")
                        ui.label("Brak wgranych dokumentów").classes(
                            "text-slate-600 dark:text-slate-400 font-semibold mt-2 text-sm"
                        )
                        ui.label(
                            "Przeciągnij pliki PDF/Word/Excel do pola powyżej lub w czacie."
                        ).classes("text-slate-500 text-xs mt-1")
                    return

                for fn, doc in docs_map.items():
                    card_border = "border-amber-300 dark:border-amber-700/60 bg-amber-50/40 dark:bg-amber-950/20" if doc["is_legal"] else "border-slate-200 dark:border-slate-700 bg-white dark:bg-slate-800"
                    with ui.card().classes(
                        f"w-full {card_border} border p-3 rounded-xl shadow-sm mb-2"
                    ):
                        with ui.row().classes("w-full justify-between items-center no-wrap"):
                            with ui.row().classes("items-center gap-2 flex-grow overflow-hidden"):
                                icon_name = "gavel" if doc["is_legal"] else "description"
                                if fn.endswith(".pdf") and not doc["is_legal"]:
                                    icon_name = "picture_as_pdf"
                                elif (fn.endswith(".docx") or fn.endswith(".doc")) and not doc["is_legal"]:
                                    icon_name = "article"
                                elif (fn.endswith(".xlsx") or fn.endswith(".csv")) and not doc["is_legal"]:
                                    icon_name = "table_chart"

                                icon_color = "amber-600" if doc["is_legal"] else "blue-600"
                                ui.icon(icon_name, color=icon_color, size="1.6rem")
                                with ui.column().classes("gap-0 overflow-hidden"):
                                    ui.label(fn).classes(
                                        "text-xs font-bold text-slate-900 dark:text-white truncate max-w-[170px]"
                                    ).tooltip(fn)
                                    with ui.row().classes("items-center gap-1"):
                                        if doc["is_legal"]:
                                            ui.badge("Baza Prawna", color="amber-8").props("rounded text-[9px] dense")
                                        ui.label(
                                            f"{doc['chunks']} fragm. · {doc['size_kb']} KB"
                                        ).classes("text-[10px] text-slate-500 dark:text-slate-400")

                            ui.button(
                                icon="delete",
                                on_click=lambda f=fn: self.delete_document_chunks(f),
                            ).props("flat round dense color=red size=sm").tooltip("Usuń dokument")

        except Exception as e:
            with self.docs_sidebar:
                ui.label(f"Błąd ładowania: {str(e)}").classes("text-xs text-red-500")

    # --- LLM Querying & Streaming Execution ---

    async def stream_ollama_response(
        self, endpoint_url: str, model_name: str, prompt: str, system_prompt: str
    ):
        url = f"{endpoint_url.rstrip('/')}/api/generate"
        payload = {
            "model": model_name,
            "prompt": prompt,
            "system": system_prompt,
            "stream": True,
            "options": {"temperature": 0.2, "top_p": 0.9},
        }
        async with httpx.AsyncClient(timeout=180.0) as client:
            async with client.stream("POST", url, json=payload) as response:
                if response.status_code != 200:
                    yield f"⚠️ Błąd serwera LLM ({response.status_code}): {await response.aread()}"
                    return

                async for line in response.aiter_lines():
                    if line:
                        try:
                            data = json.loads(line)
                            token = data.get("response", "")
                            if token:
                                yield token
                        except Exception:
                            pass

    async def stream_openrouter_response(
        self, model_name: str, prompt: str, system_prompt: str
    ):
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://pd.dreamsoft.pro",
            "X-Title": "RAE-Suite PD Knowledge Hub",
        }
        payload = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            "stream": True,
            "temperature": 0.2,
        }
        async with httpx.AsyncClient(timeout=120.0) as client:
            async with client.stream(
                "POST", CLOUD_OPENROUTER_URL, headers=headers, json=payload
            ) as response:
                if response.status_code != 200:
                    yield f"⚠️ Błąd OpenRouter ({response.status_code}): {await response.aread()}"
                    return

                async for line in response.aiter_lines():
                    if line.startswith("data: ") and line != "data: [DONE]":
                        try:
                            data = json.loads(line[6:])
                            token = data["choices"][0]["delta"].get("content", "")
                            if token:
                                yield token
                        except Exception:
                            pass

    async def handle_send_message(self, user_input_field, custom_query: Optional[str] = None):
        query = (custom_query or user_input_field.value or "").strip()
        if not query or self.is_generating:
            return

        if not custom_query:
            user_input_field.value = ""
        self.is_generating = True

        self.messages.append(
            {"role": "user", "content": query, "time": time.strftime("%H:%M")}
        )

        with self.chat_container:
            with ui.row().classes("w-full justify-end my-3"):
                with ui.card().classes(
                    "bg-blue-600 text-white p-4 rounded-2xl max-w-2xl shadow-md border-0"
                ):
                    with ui.row().classes("w-full justify-between items-center mb-1 gap-4"):
                        ui.label(self.user_email).classes(
                            "text-xs font-bold text-blue-100"
                        )
                        ui.label(time.strftime("%H:%M")).classes(
                            "text-[10px] text-blue-200"
                        )
                    ui.label(query).classes(
                        "text-sm font-medium leading-relaxed whitespace-pre-wrap"
                    )

        model_cfg = AVAILABLE_MODELS.get(
            self.selected_model_key, AVAILABLE_MODELS["node1_bielik"]
        )
        model_display_name = model_cfg["name"]

        # 1. RAG Context Retrieval from RAE Memory
        retrieved_chunks = []
        context_text = ""
        if self.use_rag:
            try:
                search_res = await self.client.search_memories(
                    query=query,
                    project=self.project,
                    k=5,
                    layers=["semantic", "episodic"],
                )
                retrieved_chunks = search_res.get("results", [])
                if retrieved_chunks:
                    context_blocks = []
                    for idx, ch in enumerate(retrieved_chunks):
                        fn = ch.get("metadata", {}).get("filename", "Dokument")
                        chunk_no = ch.get("metadata", {}).get(
                            "chunk_index", idx + 1
                        )
                        context_blocks.append(
                            f"[Dokument {idx+1}: {fn} (Część #{chunk_no})]\n{ch.get('content', '')}"
                        )
                    context_text = "\n\n---\n\n".join(context_blocks)
            except Exception as e:
                logger.error("rag_search_error", error=str(e))

        # 2. Live Web & ISAP Sejm Search if enabled
        web_results = []
        if self.use_web_search:
            try:
                web_results = await self.search_isap_and_web(query)
            except Exception as e:
                logger.warning("live_web_search_error", error=str(e))

        web_context_str = ""
        if web_results:
            web_blocks = []
            for w in web_results:
                web_blocks.append(f"[{w['source']}: {w['title']}]\n{w['snippet']}")
            web_context_str = "\n\n".join(web_blocks)

        # 3. System Prompt Assembly
        legal_instructions = ""
        if self.use_legal_audit:
            legal_instructions = (
                "TRYB AUDYTOR PRAWNY & RYZYKO UMOWNE (PRINT & DISPLAY):\n"
                "1. Przeprowadź rygorystyczny audyt prawny i wskaż potencjalne zagrożenia dla firmy Print & Display.\n"
                "2. Sprawdź 7 kluczowych filarów:\n"
                "   - [Prawa Autorskie / Licencje na Grafiki]: Kto odpowiada za prawa do materiałów i znaków towarowych przekazanych do druku?\n"
                "   - [Kary Umowne & Odpowiedzialność]: Czy kary nie są jednostronne, rażąco wygórowane (Art. 484 k.c.) i czy uwzględniają opóźnienia klienta?\n"
                "   - [Procedury Reklamacyjne & Fogra/ISO]: Czy określono tolerancje wymiarowe i akceptację proofa barwnego?\n"
                "   - [Zatory Płatnicze & Terminy B2B]: Czy terminy płatności nie przekraczają 60 dni i przewidują rekompensatę 40/70/100 EUR?\n"
                "   - [Klauzule Niedozwolone / Abuzywne]: Czy w umowie nie ma zapisów sprzecznych z ustawą lub orzecznictwem UOKiK?\n"
                "   - [Ochrona Danych / RODO]: Czy przewidziano umowę powierzenia przy materiałach personalizowanych?\n"
                "   - [Siła Wyższa & Dostawy Surowców]: Jak uregulowano brak surowca / awarie maszyn?\n"
                "3. Oznacz każde ryzyko oceną: 🔴 WYSOKIE RYZYKO, 🟡 ŚREDNIE RYZYKO, 🟢 BEZPIECZNE / ZGODNE.\n"
                "4. Zaproponuj bezpieczny zapis alternatywny (Proponowana klauzula ochronna) dla każdego wykrytego ryzyka.\n"
            )

        system_prompt = (
            f"Jesteś wyspecjalizowanym Asystentem Prawnym i Wiedzy Korporacyjnej firmy Print & Display (Projekt: pd-test-llm).\n"
            f"Odpowiadaj profesjonalnie, ściśle w języku polskim, z zachowaniem terminologii polskiego prawa i norm poligraficznych.\n\n"
            f"{legal_instructions}\n"
            f"### KONTEKST Z BAZY DOKUMENTÓW FIRMOWYCH (RAE MEMORY):\n{context_text if context_text else 'Brak bezpośrednich dokumentów w pamięci lokalnej.'}\n\n"
            f"### AKTUALNE PRZEPISY Z INTERNETU & BAZY ISAP SEJM:\n{web_context_str if web_context_str else 'Brak dodatkowych wyników online.'}\n"
        )

        # 4. Model Response Bubble
        with self.chat_container:
            with ui.row().classes("w-full items-start gap-3 my-3"):
                ui.avatar(
                    icon=model_cfg["icon"],
                    color=model_cfg["badge_color"] + "-6",
                    text_color="white",
                ).classes("shadow-md mt-1")

                with ui.card().classes(
                    "flex-grow bg-white dark:bg-slate-800 border border-slate-200 dark:border-slate-700 p-5 rounded-2xl shadow-sm max-w-3xl"
                ):
                    with ui.row().classes("w-full justify-between items-center mb-2"):
                        with ui.row().classes("items-center gap-2"):
                            ui.badge(
                                model_cfg["short_name"],
                                color=model_cfg["badge_color"] + "-7",
                            ).props("rounded dense")
                            if self.use_legal_audit:
                                ui.badge("Audyt Prawny", color="amber-8").props("rounded dense")
                            ui.label(model_cfg["device"]).classes(
                                "text-[11px] font-semibold text-slate-500 dark:text-slate-400"
                            )
                        ui.label(time.strftime("%H:%M")).classes(
                            "text-[10px] text-slate-400"
                        )

                    response_markdown = ui.markdown("").classes(
                        "text-slate-900 dark:text-slate-100 text-sm leading-relaxed"
                    )
                    spinner = ui.spinner("dots", size="md", color="blue-600")

                    # Sources Collapsible Blocks
                    if retrieved_chunks or web_results:
                        with ui.column().classes("w-full gap-2 mt-3"):
                            if retrieved_chunks:
                                with ui.expansion(
                                    f"📚 Wykorzystane źródła z bazy dokumentów ({len(retrieved_chunks)})",
                                    icon="menu_book",
                                ).classes(
                                    "w-full text-xs bg-slate-50 dark:bg-slate-900/60 rounded-xl text-slate-700 dark:text-slate-300 border border-slate-200 dark:border-slate-800"
                                ):
                                    for r in retrieved_chunks:
                                        fn = r.get("metadata", {}).get("filename", "Dokument")
                                        score = round(r.get("score", 0.0) * 100, 1)
                                        with ui.card().classes(
                                            "w-full bg-white dark:bg-slate-950 p-2.5 my-1 border border-slate-200 dark:border-slate-800 rounded-lg"
                                        ):
                                            ui.label(
                                                f"📄 {fn} (Dopasowanie: {score}%)"
                                            ).classes(
                                                "font-bold text-blue-700 dark:text-blue-300 text-xs"
                                            )
                                            ui.label(
                                                r.get("content", "")[:220] + "..."
                                            ).classes(
                                                "text-[11px] text-slate-600 dark:text-slate-400 italic"
                                            )

                            if web_results:
                                with ui.expansion(
                                    f"🌐 Wyniki z bazy prawnej ISAP Sejm RP & Internetu ({len(web_results)})",
                                    icon="public",
                                ).classes(
                                    "w-full text-xs bg-blue-50/50 dark:bg-slate-900/60 rounded-xl text-slate-700 dark:text-slate-300 border border-blue-200 dark:border-slate-800"
                                ):
                                    for w in web_results:
                                        with ui.card().classes(
                                            "w-full bg-white dark:bg-slate-950 p-2.5 my-1 border border-slate-200 dark:border-slate-800 rounded-lg"
                                        ):
                                            ui.label(f"🏛️ {w['source']}: {w['title']}").classes(
                                                "font-bold text-indigo-700 dark:text-indigo-300 text-xs"
                                            )
                                            ui.label(w['snippet']).classes(
                                                "text-[11px] text-slate-600 dark:text-slate-400"
                                            )

        full_answer = ""
        try:
            if model_cfg["type"] == "ollama":
                stream_gen = self.stream_ollama_response(
                    endpoint_url=model_cfg["url"],
                    model_name=model_cfg["model"],
                    prompt=query,
                    system_prompt=system_prompt,
                )
            else:
                stream_gen = self.stream_openrouter_response(
                    model_name=model_cfg["model"],
                    prompt=query,
                    system_prompt=system_prompt,
                )

            async for token in stream_gen:
                full_answer += token
                response_markdown.set_content(full_answer)

            spinner.set_visibility(False)

            # Store turn in RAE Memory (Episodic layer)
            await self.client.store_memory(
                content=f"Pytanie ({self.user_email}): {query}\nOdpowiedź ({model_display_name}): {full_answer}",
                project=self.project,
                source=f"chat:{self.selected_model_key}",
                tags=["pd_chat_history", self.user_email, "legal_audit" if self.use_legal_audit else "general"],
                layer="episodic",
                metadata={
                    "user": self.user_email,
                    "model": model_display_name,
                    "rag_used": bool(retrieved_chunks),
                    "web_used": bool(web_results),
                    "legal_audit": self.use_legal_audit,
                },
            )

        except Exception as stream_err:
            spinner.set_visibility(False)
            response_markdown.set_content(
                f"{full_answer}\n\n⚠️ **Błąd połączenia z modelem:** {str(stream_err)}"
            )

        self.is_generating = False

    def select_model(self, key: str):
        self.selected_model_key = key
        cfg = AVAILABLE_MODELS.get(key, AVAILABLE_MODELS["node1_bielik"])
        if self.model_select_widget:
            self.model_select_widget.value = key
        if self.active_model_badge:
            self.active_model_badge.text = cfg["short_name"]
            self.active_model_badge.props(f"color={cfg['badge_color']}-7")
        ui.notify(
            f"Przełączono model na: {cfg['name']}",
            type="positive",
            position="top-right",
            duration=2.5,
        )

    def trigger_quick_audit(self, audit_type: str):
        prompts = {
            "full_audit": "Przeprowadź kompleksowy audyt prawny i analizę ryzyka dla wgranych dokumentów i umów Print & Display pod kątem polskiego prawa (Kodeks Cywilny, Kary Umowne, Prawa Autorskie, Zatory Płatnicze, Fogra/ISO).",
            "copyright": "Zbadaj zgodność zapisów dotyczących praw autorskich do projektów graficznych i znaków towarowych z Ustawą o prawie autorskim (art. 41-68). Czy drukarnia jest zabezpieczona przed roszczeniami osób trzecich?",
            "penalties": "Przeanalizuj kary umowne i terminy płatności w umowach. Czy kary są symetryczne, czy nie przekraczają granic swobody umów (art. 484 k.c.) oraz czy terminy są zgodne z Ustawą o zatorach płatniczych?",
            "quality": "Sprawdź zapisy dotyczące odbiorów jakościowych, reklamacji i tolerancji produkcyjnych (standardy Fogra 39/51, tolerancje cięcia Kongsberg). Jakie ryzyka rodzi brak fizycznego proofa barwnego?",
        }
        query = prompts.get(audit_type, prompts["full_audit"])
        if self.prompt_input:
            asyncio.create_task(self.handle_send_message(self.prompt_input, custom_query=query))

    def toggle_theme(self):
        if self.dark_mode_controller:
            self.is_dark_mode = not self.is_dark_mode
            self.dark_mode_controller.value = self.is_dark_mode

    # --- Render Main View ---

    def render(self, standalone: bool = False):
        self.dark_mode_controller = ui.dark_mode(value=False)

        # Base Page Container
        with ui.column().classes(
            "w-full min-h-screen bg-slate-50 dark:bg-slate-950 text-slate-900 dark:text-slate-100 p-3 md:p-6 transition-colors duration-200"
        ):
            # 1. TOP HEADER BAR
            with ui.card().classes(
                "w-full bg-white dark:bg-slate-900 border border-slate-200 dark:border-slate-800 px-6 py-4 rounded-2xl shadow-sm mb-4"
            ):
                with ui.row().classes(
                    "w-full justify-between items-center gap-4 flex-wrap"
                ):
                    with ui.row().classes("items-center gap-3"):
                        ui.icon("gavel", color="blue-600", size="2.4rem")
                        with ui.column().classes("gap-0"):
                            ui.label("Print & Display AI Knowledge & Legal Hub").classes(
                                "text-xl md:text-2xl font-black tracking-tight text-slate-900 dark:text-white"
                            )
                            ui.label(
                                "Analiza dokumentów, audyt zgodności z polskim prawem (ISAP Sejm) i czat wielomodelowy (Projekt: pd-test-llm)"
                            ).classes("text-xs text-slate-500 dark:text-slate-400")

                    with ui.row().classes("items-center gap-2 flex-wrap"):
                        ui.badge("pd-test-llm", color="blue-7").props(
                            "rounded"
                        ).tooltip("Projekt RAE Memory")

                        ui.badge(
                            f"👤 {self.user_email}",
                            color="slate-200 dark:slate-800",
                        ).props("outline text-color=slate-800 dark:text-color=white")

                        ui.button(
                            "🌐 RAE Suite Portal",
                            icon="open_in_new",
                            on_click=lambda: ui.navigate.to("https://rae.dreamsoft.pro/", new_tab=True),
                        ).props("outline dense color=blue-7 size=sm rounded").tooltip(
                            "Przejdź do RAE Suite Portal (https://rae.dreamsoft.pro/)"
                        )

                        ui.button(
                            icon="dark_mode",
                            on_click=self.toggle_theme,
                        ).props(
                            "flat round dense color=slate-700 dark:color=amber"
                        ).tooltip("Przełącz motyw Jasny / Ciemny")

                        if standalone:
                            ui.button(
                                "Wyloguj",
                                icon="logout",
                                on_click=lambda: ui.navigate.to("/logout"),
                            ).props("flat dense color=red size=sm")

            # 2. PROMINENT MODEL SELECTOR & QUICK AUDIT TOOLBAR
            with ui.card().classes(
                "w-full bg-white dark:bg-slate-900 border border-slate-200 dark:border-slate-800 p-4 rounded-2xl shadow-sm mb-4"
            ):
                with ui.column().classes("w-full gap-3"):
                    # Row 1: Model Selection
                    with ui.row().classes("w-full justify-between items-center gap-4 flex-wrap"):
                        with ui.row().classes("items-center gap-2"):
                            ui.icon("memory", color="blue-600", size="1.6rem")
                            ui.label("Aktywny Model LLM:").classes(
                                "text-sm font-bold text-slate-800 dark:text-slate-200"
                            )
                            initial_cfg = AVAILABLE_MODELS[self.selected_model_key]
                            self.active_model_badge = ui.badge(
                                initial_cfg["short_name"],
                                color=initial_cfg["badge_color"] + "-7",
                            ).props("rounded text-xs")

                        with ui.row().classes("items-center gap-2 flex-wrap"):
                            ui.label("Szybki wybór:").classes("text-xs text-slate-400 font-medium")

                            ui.button(
                                "🚀 Bielik 11B (Node 1 GPU)",
                                on_click=lambda: self.select_model("node1_bielik"),
                            ).props("outline dense size=sm color=indigo rounded").tooltip(
                                "Polski model SpeakLeash na karcie NVIDIA RTX 4080 SUPER (Node 1) - Znakomity do polskiego prawa"
                            )

                            ui.button(
                                "🧠 DeepSeek-R1 (Node 1)",
                                on_click=lambda: self.select_model("node1_deepseek_r1"),
                            ).props("outline dense size=sm color=purple rounded").tooltip(
                                "Model wnioskujący Reasoning - Doskonały do wykrywania luk w umowach"
                            )

                            ui.button(
                                "⚡ Qwen 3.5 9B (Node 1)",
                                on_click=lambda: self.select_model("node1_qwen35"),
                            ).props("outline dense size=sm color=blue rounded")

                            ui.button(
                                "💻 Bielik (Laptop)",
                                on_click=lambda: self.select_model("laptop_bielik"),
                            ).props("outline dense size=sm color=teal rounded")

                            self.model_select_widget = ui.select(
                                options={k: v["name"] for k, v in AVAILABLE_MODELS.items()},
                                value=self.selected_model_key,
                                label="Wszystkie modele & urządzenia",
                            ).classes("w-72").props("dense outlined")
                            self.model_select_widget.on(
                                "update:model_value",
                                lambda e: self.select_model(e.value),
                            )

                    ui.separator().classes("bg-slate-100 dark:bg-slate-800")

                    # Row 2: 1-Click Legal Audits & Toggles
                    with ui.row().classes("w-full justify-between items-center gap-3 flex-wrap"):
                        with ui.row().classes("items-center gap-2 flex-wrap"):
                            ui.label("⚡ Szybki Audyt Prawny:").classes(
                                "text-xs font-bold text-amber-700 dark:text-amber-400"
                            )
                            ui.button(
                                "⚖️ Pełny Audyt Umowy",
                                on_click=lambda: self.trigger_quick_audit("full_audit"),
                            ).props("unelevated dense size=sm color=amber-8 rounded").tooltip(
                                "Skanuje dokument pod kątem 7 kluczowych ryzyk prawnych"
                            )
                            ui.button(
                                "🎨 Prawa Autorskie & Licencje",
                                on_click=lambda: self.trigger_quick_audit("copyright"),
                            ).props("outline dense size=sm color=amber-9 rounded")
                            ui.button(
                                "💰 Kary Umowne & Płatności",
                                on_click=lambda: self.trigger_quick_audit("penalties"),
                            ).props("outline dense size=sm color=amber-9 rounded")
                            ui.button(
                                "📐 Reklamacje & Fogra/ISO",
                                on_click=lambda: self.trigger_quick_audit("quality"),
                            ).props("outline dense size=sm color=amber-9 rounded")

                        with ui.row().classes("items-center gap-4 text-xs"):
                            legal_switch = ui.switch(
                                "⚖️ Tryb Audytora Prawnego",
                                value=self.use_legal_audit,
                            ).props("dense color=amber")
                            legal_switch.on(
                                "update:model_value",
                                lambda e: setattr(self, "use_legal_audit", e.value),
                            )

                            web_switch = ui.switch(
                                "🌐 Szukaj w ISAP Sejm & Sieci",
                                value=self.use_web_search,
                            ).props("dense color=blue")
                            web_switch.on(
                                "update:model_value",
                                lambda e: setattr(self, "use_web_search", e.value),
                            )

            # 3. MAIN WORKSPACE: 2 COLUMNS (CHAT + DOCUMENTS SIDEBAR)
            with ui.row().classes("w-full gap-4 items-stretch"):
                # LEFT MAIN COLUMN: CHAT & PROMPT INPUT
                with ui.column().classes("flex-grow min-w-[320px] gap-4"):
                    # Chat Window Card
                    with ui.card().classes(
                        "w-full bg-white dark:bg-slate-900 border border-slate-200 dark:border-slate-800 p-5 rounded-2xl shadow-sm flex flex-col justify-between min-h-[560px]"
                    ):
                        with ui.scroll_area().classes("w-full flex-grow pr-2 min-h-[440px] max-h-[620px]"):
                            self.chat_container = ui.column().classes("w-full gap-3")
                            with self.chat_container:
                                with ui.row().classes("w-full items-start gap-3 my-2"):
                                    ui.avatar(icon="gavel", color="blue-7", text_color="white")
                                    with ui.card().classes(
                                        "bg-slate-50 dark:bg-slate-800/90 border border-slate-200 dark:border-slate-700 p-4 rounded-2xl max-w-2xl"
                                    ):
                                        ui.label("Asystent Wiedzy & Audytu Prawnego Print & Display").classes(
                                            "text-xs font-bold text-blue-600 dark:text-blue-400 mb-1"
                                        )
                                        ui.markdown(
                                            "Witaj! Jestem asystentem wiedzy korporacyjnej i audytu prawnego Print & Display.\n\n"
                                            "- 📎 **Wgraj dokumenty lub umowy**: kliknij spinacz poniżej lub przeciągnij plik PDF / Word do panelu po prawej.\n"
                                            "- ⚖️ **Zbadaj zgodność z polskim prawem**: kliknij przycisk *Pełny Audyt Umowy* na górnym pasku lub wpisz własne pytanie.\n"
                                            "- 🏛️ **Baza ISAP & Ustawy**: asystent automatycznie weryfikuje Kodeks Cywilny, Prawo Autorskie, Zatory Płatnicze i bazy Sejmu RP.\n"
                                            "- 🚀 **Wybierz model**: do analizy polskiego prawa polecamy **Bielik 11B (SpeakLeash)** lub **DeepSeek-R1 (Node 1)**."
                                        ).classes("text-slate-800 dark:text-slate-200 text-sm leading-relaxed")

                        # IN-CHAT PROMPT & UPLOAD BAR
                        with ui.column().classes("w-full pt-4 border-t border-slate-200 dark:border-slate-800 gap-2"):
                            with ui.row().classes("w-full justify-between items-center text-xs"):
                                rag_switch = ui.switch(
                                    "🔍 Szukaj w dokumentach (RAG)",
                                    value=self.use_rag,
                                ).props("dense color=blue")
                                rag_switch.on(
                                    "update:model_value",
                                    lambda e: setattr(self, "use_rag", e.value),
                                )

                                ui.button(
                                    "Wyczyść rozmowę",
                                    icon="delete_sweep",
                                    on_click=lambda: (
                                        self.chat_container.clear(),
                                        self.messages.clear(),
                                    ),
                                ).props("flat dense color=slate-500 size=xs")

                            # Prompt Row
                            with ui.row().classes("w-full items-center gap-2"):
                                upload_widget = ui.upload(
                                    on_upload=self.handle_file_upload,
                                    max_file_size=50_000_000,
                                    auto_upload=True,
                                    multiple=True,
                                ).props(
                                    'accept=".pdf,.docx,.doc,.txt,.md,.csv,.xlsx,.json" flat dense'
                                ).classes("hidden")

                                ui.button(
                                    icon="attach_file",
                                    on_click=upload_widget.run_method("pickFiles"),
                                ).props(
                                    "flat round color=blue-7 size=md"
                                ).tooltip("Wgraj dokument (PDF, Word, Excel, TXT) do bazy wiedzy")

                                self.prompt_input = ui.input(
                                    placeholder="Napisz pytanie do dokumentów lub wklej treść umowy do audytu prawnego... (Wciśnij Enter aby wysłać)",
                                ).classes(
                                    "flex-grow bg-slate-50 dark:bg-slate-900 text-slate-900 dark:text-white rounded-xl"
                                ).props(
                                    "outlined dense input-class='text-slate-900 dark:text-white text-sm font-medium'"
                                )

                                self.prompt_input.on(
                                    "keydown.enter",
                                    lambda: self.handle_send_message(self.prompt_input),
                                )

                                ui.button(
                                    icon="send",
                                    on_click=lambda: self.handle_send_message(self.prompt_input),
                                ).props("elevated color=blue-7 rounded-xl size=md").tooltip(
                                    "Wyślij zapytanie"
                                )

                # RIGHT SIDEBAR: DOCUMENTS & LEGAL CORPUS
                with ui.card().classes(
                    "w-full lg:w-80 bg-white dark:bg-slate-900 border border-slate-200 dark:border-slate-800 p-4 rounded-2xl shadow-sm flex flex-col justify-between"
                ):
                    with ui.column().classes("w-full gap-3"):
                        with ui.row().classes("w-full justify-between items-center pb-2 border-b border-slate-100 dark:border-slate-800"):
                            with ui.row().classes("items-center gap-2"):
                                ui.icon("folder_special", color="blue-600", size="1.4rem")
                                ui.label("Baza Dokumentów & Prawa").classes(
                                    "text-sm font-bold text-slate-900 dark:text-white"
                                )
                            self.docs_badge_label = ui.badge("0 dok.", color="blue-7").props("rounded dense")

                        # Drag & Drop Zone
                        ui.upload(
                            on_upload=self.handle_file_upload,
                            max_file_size=50_000_000,
                            auto_upload=True,
                            multiple=True,
                        ).props(
                            'accept=".pdf,.docx,.doc,.txt,.md,.csv,.xlsx,.json" label="Przeciągnij pliki tutaj 📂"'
                        ).classes(
                            "w-full border-2 border-dashed border-slate-300 dark:border-slate-700 bg-slate-50 dark:bg-slate-950/60 rounded-xl"
                        )

                        # Documents list
                        with ui.scroll_area().classes("w-full max-h-[380px] pr-1"):
                            self.docs_sidebar = ui.column().classes("w-full")

                    with ui.column().classes("w-full pt-3 border-t border-slate-100 dark:border-slate-800"):
                        ui.button(
                            "Odśwież listę dokumentów",
                            icon="refresh",
                            on_click=self.refresh_documents_list,
                        ).props("flat dense color=blue-7 size=sm w-full")

            # Fire initial document load
            ui.timer(0.3, self.refresh_documents_list, once=True)
